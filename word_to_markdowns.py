from docx import Document
import os


def word_to_markdowns(word_file, output_dir):
    # 打开Word文档
    doc = Document(word_file)

    # 遍历文档中的所有段落
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == '文章标题':  # 假设“标题”样式对应文章标题
            article_title = paragraph.text.strip()
            with open(os.path.join(output_dir, f"{article_title}.md"), 'w', encoding='utf-8') as f:
                f.write(f"# {article_title}\n\n")  # Markdown文件开头写入标题

                # 寻找下一个标题之前的所有内容作为文章正文
                content_paragraphs = []
                for next_p in doc.paragraphs[i + 1:]:
                    if next_p.style.name == '文章标题':
                        break
                    else:
                        content_paragraphs.append(next_p)

                # 将正文段落转为Markdown格式并写入文件
                for p in content_paragraphs:
                    text = p.text.strip()
                    f.write(text + '\n\n')


# 使用方法
# word_to_markdowns(r'C:\Users\saber\Desktop\pythonscript\xxck\学习参考2022年第15期-中国共产党成立一百零一周年.docx',
#                   r'C:\Users\saber\Desktop\pythonscript\xxck')

# 一个文件夹有多个word文件，一个Word文件里面有多篇文章，每篇文章都有自己的标题，把每个word文件转成一个markdown文件，文件名就是word文件名，
# 在markdown里面是一级标题，下面两行分别是[[toc]]和[toc],在遇到第一个文章标题前的内容都忽略
# 其他具体的格式对应为：
# 1.word里面样式名为文章标题的，在markdown文件中使用二级标题。
# 2.word里面样式名为黑体小标题或黑体小标题正的，在markdown文件中使用加粗。
# 3.word里面样式名为文章一级标题的，在markdown文件中使用三级标题。
# 4.word里面样式名为文章二级标题的，在markdown文件中使用四级标题。
# 5.其他的就是正文，在markdown文件中使用普通文本。

# 指定word文件夹路径和输出markdown文件夹路径
word_folder_path = r'C:\Users\saber\Desktop\pythonscript\xxck'
output_folder_path = r'C:\Users\saber\Desktop\pythonscript\xxck\111'

if not os.path.exists(output_folder_path):
    os.makedirs(output_folder_path)


def convert_word_to_markdown(word_file_path):
    # 读取word文件
    doc = Document(word_file_path)

    # 获取word文件名作为markdown文件名
    md_file_name = os.path.splitext(os.path.basename(word_file_path))[0] + ".md"
    md_file_path = os.path.join(output_folder_path, md_file_name)

    # 初始化markdown内容并忽略第一个文章标题前的内容
    article_started = False
    md_content = f'# {os.path.splitext(os.path.basename(word_file_path))[0]}\n\n[[toc]]\n\n[toc]\n\n'

    for para in doc.paragraphs:
        style = para.style
        if hasattr(style, 'name'):  # 假设style对象有name属性，实际情况请查阅style对象或根据style_id判断
            if not article_started and style.name == '文章标题':  # 遇到第一个文章标题后开始记录内容
                article_started = True
                md_content += '## ' + para.text + '\n\n'
            elif article_started:
                if style.name == '文章标题':
                    md_content += '## ' + para.text + '\n\n'
                elif style.name in ['黑体小标题', '黑体小标题正']:
                    md_content += '**' + para.text + '**\n\n'
                elif style.name == '文章一级标题':
                    md_content += '### ' + para.text + '\n\n'
                elif style.name == '文章二级标题':
                    md_content += '#### ' + para.text + '\n\n'
                elif article_started:  # 其他正文样式（在遇到第一个文章标题后）
                    md_content += para.text + '\n\n'

    # 将markdown内容写入文件
    with open(md_file_path, 'w', encoding='utf-8') as f:
        f.write(md_content)


# 遍历word文件夹中的所有word文件
for file in os.listdir(word_folder_path):
    if file.endswith(".docx"):
        word_file_path = os.path.join(word_folder_path, file)
        convert_word_to_markdown(word_file_path)

# 注意：以上代码仅提供了一个基本的结构框架，实际操作时需要处理样式映射的具体逻辑。
